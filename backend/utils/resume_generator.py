from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import json
from datetime import datetime

def generate_resume_docx(resume_data):
    """
    生成Word格式的简历
    
    Args:
        resume_data: dict包含简历信息
            - name: 姓名
            - email: 邮箱
            - phone: 电话
            - education: 教育经历列表
            - work_experience: 工作经历列表
            - skills: 技能列表
            - self_introduction: 自我介绍
    
    Returns:
        str: 生成的文件名
    """
    doc = Document()
    
    # 设置页边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # 添加标题 - 姓名
    title = doc.add_heading(resume_data['name'], 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # 添加联系方式
    contact = doc.add_paragraph()
    contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact_text = f"{resume_data['phone']} | {resume_data['email']}"
    contact_run = contact.add_run(contact_text)
    contact_run.font.size = Pt(11)
    
    doc.add_paragraph()  # 空行
    
    # 自我介绍
    if resume_data.get('self_introduction'):
        doc.add_heading('自我介绍', level=1)
        intro = doc.add_paragraph(resume_data['self_introduction'])
        intro.style.font.size = Pt(11)
        doc.add_paragraph()
    
    # 教育经历
    if resume_data.get('education'):
        doc.add_heading('教育经历', level=1)
        education_list = resume_data['education']
        if isinstance(education_list, str):
            education_list = json.loads(education_list)
        
        for edu in education_list:
            p = doc.add_paragraph(style='List Bullet')
            
            # 学校和专业
            school_run = p.add_run(f"{edu.get('school', '')} - {edu.get('major', '')}")
            school_run.bold = True
            school_run.font.size = Pt(11)
            
            # 时间和学历
            if edu.get('start_date') or edu.get('end_date') or edu.get('degree'):
                p.add_run('\n')
                time_text = f"{edu.get('start_date', '')} - {edu.get('end_date', '')}"
                if edu.get('degree'):
                    time_text += f" | {edu.get('degree')}"
                time_run = p.add_run(time_text)
                time_run.font.size = Pt(10)
                time_run.font.color.rgb = RGBColor(96, 96, 96)
            
            if edu.get('description'):
                p.add_run(f"\n{edu.get('description')}")
        
        doc.add_paragraph()
    
    # 工作经历
    if resume_data.get('work_experience'):
        doc.add_heading('工作经历', level=1)
        experience_list = resume_data['work_experience']
        if isinstance(experience_list, str):
            experience_list = json.loads(experience_list)
        
        for exp in experience_list:
            p = doc.add_paragraph(style='List Bullet')
            
            # 公司和职位
            company_run = p.add_run(f"{exp.get('company', '')} - {exp.get('position', '')}")
            company_run.bold = True
            company_run.font.size = Pt(11)
            
            # 时间
            if exp.get('start_date') or exp.get('end_date'):
                p.add_run('\n')
                time_text = f"{exp.get('start_date', '')} - {exp.get('end_date', '')}"
                time_run = p.add_run(time_text)
                time_run.font.size = Pt(10)
                time_run.font.color.rgb = RGBColor(96, 96, 96)
            
            # 工作描述
            if exp.get('description'):
                p.add_run(f"\n{exp.get('description')}")
        
        doc.add_paragraph()
    
    # 技能
    if resume_data.get('skills'):
        doc.add_heading('专业技能', level=1)
        skills_list = resume_data['skills']
        if isinstance(skills_list, str):
            skills_list = json.loads(skills_list)
        
        if isinstance(skills_list, list):
            for skill in skills_list:
                if isinstance(skill, dict):
                    skill_text = skill.get('name', '')
                    if skill.get('level'):
                        skill_text += f" - {skill.get('level')}"
                else:
                    skill_text = str(skill)
                
                doc.add_paragraph(skill_text, style='List Bullet')
        else:
            doc.add_paragraph(str(skills_list))
    
    # 保存文件
    filename = f"resume_{resume_data['name']}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 
                             '..', 'uploads', 'generated')
    os.makedirs(output_dir, exist_ok=True)
    
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    
    return filename
